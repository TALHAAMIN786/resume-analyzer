"""
parser.py — Resume file parser for PDF and DOCX
Phase 2: Extract clean text from uploaded resume files
"""

import re
import pdfplumber
from docx import Document
from config import ALLOWED_EXTENSIONS


# ══════════════════════════════════════════════════════════════════════════════
#  TEXT CLEANING
# ══════════════════════════════════════════════════════════════════════════════

def clean_text(text: str) -> str:
    """
    Clean raw extracted text:
    - Remove excessive whitespace and blank lines
    - Fix broken line spacing
    - Strip non-printable characters
    """
    if not text:
        return ""

    # Remove non-printable characters (keep newlines and tabs)
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E]', ' ', text)

    # Replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)

    # Replace 3+ newlines with exactly 2 (preserve section breaks)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Strip trailing whitespace from each line
    lines = [line.rstrip() for line in text.split('\n')]

    # Remove lines that are just spaces or single characters (noise)
    lines = [line for line in lines if len(line.strip()) > 1 or line == '']

    return '\n'.join(lines).strip()


# ══════════════════════════════════════════════════════════════════════════════
#  PDF PARSER
# ══════════════════════════════════════════════════════════════════════════════

def parse_pdf(file_path: str) -> dict:
    """
    Extract text from a PDF resume using pdfplumber.
    Returns a dict with: text, page_count, success, error
    """
    result = {
        "text": "",
        "page_count": 0,
        "success": False,
        "error": None
    }

    try:
        with pdfplumber.open(file_path) as pdf:
            result["page_count"] = len(pdf.pages)
            all_text = []

            for i, page in enumerate(pdf.pages):
                # Extract text from page
                page_text = page.extract_text()

                if page_text:
                    all_text.append(page_text)
                else:
                    # Try extracting words if extract_text fails
                    words = page.extract_words()
                    if words:
                        word_text = ' '.join([w['text'] for w in words])
                        all_text.append(word_text)

            raw_text = '\n\n'.join(all_text)
            result["text"]    = clean_text(raw_text)
            result["success"] = bool(result["text"])

            if not result["text"]:
                result["error"] = "No readable text found. PDF may be image-based or scanned."

    except Exception as e:
        result["error"] = f"PDF parsing failed: {str(e)}"

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX PARSER
# ══════════════════════════════════════════════════════════════════════════════

def parse_docx(file_path: str) -> dict:
    """
    Extract text from a DOCX resume using python-docx.
    Returns a dict with: text, page_count, success, error
    """
    result = {
        "text": "",
        "page_count": 1,   # DOCX doesn't expose page count easily
        "success": False,
        "error": None
    }

    try:
        doc = Document(file_path)
        all_text = []

        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                all_text.append(para_text)

        # Also extract text from tables (skills tables are common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    all_text.append(row_text)

        raw_text = '\n'.join(all_text)
        result["text"]    = clean_text(raw_text)
        result["success"] = bool(result["text"])

        if not result["text"]:
            result["error"] = "No readable text found in DOCX file."

    except Exception as e:
        result["error"] = f"DOCX parsing failed: {str(e)}"

    return result


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION DETECTOR
# ══════════════════════════════════════════════════════════════════════════════

# Common resume section heading keywords
SECTION_KEYWORDS = {
    "contact":    ["contact", "email", "phone", "linkedin", "github", "address"],
    "summary":    ["summary", "objective", "profile", "about", "overview"],
    "experience": ["experience", "employment", "work history", "career", "positions"],
    "education":  ["education", "academic", "degree", "university", "college", "school"],
    "skills":     ["skills", "technologies", "tools", "competencies", "expertise"],
    "projects":   ["projects", "portfolio", "works", "achievements"],
    "certifications": ["certifications", "certificates", "licenses", "credentials"],
    "awards":     ["awards", "honors", "recognition", "achievements"],
}

def detect_sections(text: str) -> dict:
    """
    Scan resume text and identify which standard sections are present.
    Returns dict of section_name -> True/False
    """
    text_lower = text.lower()
    found = {}

    for section, keywords in SECTION_KEYWORDS.items():
        found[section] = any(kw in text_lower for kw in keywords)

    return found


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def parse_resume(file_path: str) -> dict:
    """
    Master function — detects file type and parses accordingly.
    Returns full result dict with text, sections, metadata.
    """
    ext = file_path.split(".")[-1].lower()

    if ext == "pdf":
        result = parse_pdf(file_path)
    elif ext == "docx":
        result = parse_docx(file_path)
    else:
        return {
            "text": "",
            "page_count": 0,
            "success": False,
            "error": f"Unsupported file type: .{ext}",
            "sections": {},
            "word_count": 0
        }

    # Add section detection and word count
    if result["success"]:
        result["sections"]   = detect_sections(result["text"])
        result["word_count"] = len(result["text"].split())
    else:
        result["sections"]   = {}
        result["word_count"] = 0

    return result


# ── Quick test (run this file directly) ───────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        path = sys.argv[1]
        print(f"Parsing: {path}")
        r = parse_resume(path)
        print(f"Success    : {r['success']}")
        print(f"Pages      : {r['page_count']}")
        print(f"Word count : {r['word_count']}")
        print(f"Sections   : {r['sections']}")
        print(f"\n--- First 500 chars ---\n{r['text'][:500]}")
    else:
        print("Usage: python parser.py path/to/resume.pdf")